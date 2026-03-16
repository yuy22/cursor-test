# -*- coding: utf-8 -*-
"""
俞正强《种子课》DOCX -> Markdown 数据清洗转换脚本
数学符号均为简单 Unicode，直接保留，无需 LaTeX 转换
"""
import sys, re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from pathlib import Path

INPUT  = r"C:\Users\b886855456ly\Desktop\俞正强种子课      一个数学特级教师的思与行.docx"
OUTPUT = r"C:\Users\b886855456ly\Desktop\俞正强种子课_cleaned.md"

# ──────────────────────────────────────────────
# 1. 噪声过滤规则（返回 True 表示该行应丢弃）
# ──────────────────────────────────────────────
NOISE_PATTERNS = [
    # 水印行（多种残缺变体）
    re.compile(r'[仅供个人科研教学使用！!]{4,}'),
    re.compile(r'供个人科研教学'),
    re.compile(r'人科研教学使用'),
    # 纯页码行：001 / 002 / 1 等
    re.compile(r'^\d{1,4}$'),
    # 拼音行（全大写英文，书名拼音）
    re.compile(r'^[A-Z\s]{10,}$'),
    # 出版/版权行
    re.compile(r'^(出版人|策划编辑|责任编辑|责任美编|封面设计|版式设计|责任校对|责任印制|项目统筹)'),
    re.compile(r'^(ISBN|CIP|图书在版编目|中国版本图书馆)'),
    re.compile(r'^(定价|传\s*真|邮\s*编|网\s*址|市场部电话|编辑部电话|社\s*址)'),
    re.compile(r'^(出版发行|印\s*刷|经\s*销|制\s*作|开\s*本|印\s*张|字\s*数|印\s*次|版\s*次|印\s*数)'),
    re.compile(r'(教育科学出版社|北京师范大学出版|人民教育出版社)'),
    # CIP 分类号行  I.① II.① 等
    re.compile(r'^[IⅠⅡⅢ\u2160-\u2169][\.\s·①②]'),
    # 格式装饰行：· 北 京 · 等（单纯地理/装饰，无实意）
    re.compile(r'^[·\s]+[\u4e00-\u9fff\s]+[·\s]+$'),
    # 孤立人名行（1-4个字被空格分隔，无上下文）
    re.compile(r'^[\u4e00-\u9fff]\s{1,3}[\u4e00-\u9fff]\s{1,3}[\u4e00-\u9fff]?\s*$'),
    # 如有印装质量问题... 等免责声明
    re.compile(r'如有印装质量问题'),
    re.compile(r'到所购图书销售'),
    # 拼音音节行（如 ZHONGZI KE ...）
    re.compile(r'^([A-Z]{2,}\s+){3,}'),
    # 教育家书院丛书系列行
    re.compile(r'教育家[书書]院丛书'),
    # OCR合并后的出版团队行（多个职位粘连）
    re.compile(r'(出版人|策划编辑).*(责任编辑|项目统筹)'),
    # 孤立人名行：纯汉字+空格，有内部空格，总长≤15（如"刘 灿"/"何 薇 郑 莉"）
    # 合法标题"丛书序""写在前面"无内部空格，不会被匹配
    re.compile(r'^[\u4e00-\u9fff]{1,5}(\s+[\u4e00-\u9fff]{1,5}){1,5}\s*$'),
]

# 出版社信息表格关键词（第一格包含这些则跳过整个表格）
PUBLISHER_TABLE_KEYWORDS = {'出版发行', '社址', '社   址', '出版 发行', '印刷', '经销', '制作'}

def is_noise(text: str) -> bool:
    t = text.strip()
    if not t:
        return True
    for pat in NOISE_PATTERNS:
        if pat.search(t):
            return True
    return False


# ──────────────────────────────────────────────
# 2. 文本修复规则
# ──────────────────────────────────────────────
def fix_text(text: str) -> str:
    t = text.strip()

    # 去掉尾部页码 "写在前面 |001" -> "写在前面"
    t = re.sub(r'\s*[\|｜]\s*\d{1,4}\s*$', '', t)

    # 修复 OCR 字符间多余空格（纯中文段，字与字之间多余空格）
    # 规则：若一段内中文字符占比>60%且有规律的单字间空格，则压缩
    if _is_spaced_chinese(t):
        t = re.sub(r'(?<=[\u4e00-\u9fff，。！？、『』【】…])\s+(?=[\u4e00-\u9fff，。！？、『』【】…])', '', t)

    # 压缩多余空格（保留单个空格）
    t = re.sub(r'[ \t]{2,}', ' ', t)

    # 规范省略号 "……" 保留
    return t


def _is_spaced_chinese(text: str) -> bool:
    """判断是否是 OCR 导致的汉字间多余空格"""
    chinese = re.findall(r'[\u4e00-\u9fff]', text)
    if len(chinese) < 4:
        return False
    ratio = len(chinese) / max(len(text.replace(' ', '')), 1)
    # 检测是否有 "字 字" 这种单字间隔模式
    spaced = re.findall(r'[\u4e00-\u9fff] [\u4e00-\u9fff]', text)
    return ratio > 0.5 and len(spaced) >= 2


# ──────────────────────────────────────────────
# 3. 标题检测
# ──────────────────────────────────────────────
# 一级标题：书的大章（目录级别）
H1_PAT = re.compile(r'^[一二三四五六七八九十]+[、.．]\s*.{2,30}$')
# 二级标题：文章标题（通常独立一行，较短，无标点结尾）
H2_SHORT = re.compile(r'^[·*＊\*]\s*(.+)$')    # 带星号/点的导语标题
# 三级标题：带 (一) (1) 等
H3_PAT = re.compile(r'^[（\(][一二三四五六七八九\d]+[）\)]\s*.+$')
# 数字序号段落（正文列表）
ORDERED_LIST = re.compile(r'^[①②③④⑤⑥⑦⑧⑨⑩\d]+[.．、]\s*.+')

def detect_structure(text: str, prev_empty: bool) -> tuple[str, str]:
    """
    返回 (md_prefix, cleaned_text)
    prev_empty: 上一行是否为空行（用于判断段落开始）
    """
    t = text.strip()

    if H1_PAT.match(t):
        return '## ', t

    m = H2_SHORT.match(t)
    if m:
        inner = m.group(1).strip()
        # 过滤装饰性短文本（如"北 京 ·"，需至少2个有实意的中文字）
        cn_count = len(re.findall(r'[\u4e00-\u9fff]', inner))
        if cn_count >= 2:
            return '### ', inner

    if H3_PAT.match(t):
        return '#### ', t

    # 文章副标题（"—— 以xxx为例"）
    if t.startswith('——') or t.startswith('— —'):
        return '> ', t

    # 引用/题记
    if t.startswith('—') and len(t) < 40:
        return '> ', t

    return '', t


# ──────────────────────────────────────────────
# 4. 表格转 Markdown
# ──────────────────────────────────────────────
def table_to_md(table) -> str:
    rows = []
    for row in table.rows:
        cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
        rows.append(cells)

    if not rows:
        return ''

    # 去掉全空行
    rows = [r for r in rows if any(c for c in r)]
    if not rows:
        return ''

    # 跳过出版社信息表格（第一格含出版社关键词）
    first_cell = rows[0][0] if rows[0] else ''
    if any(kw in first_cell for kw in PUBLISHER_TABLE_KEYWORDS):
        return ''
    # 跳过空白结构表格（有内容的格<3个）
    all_text_cells = [c for r in rows for c in r if c.strip()]
    if len(all_text_cells) < 3:
        return ''

    # 合并重复单元格（DOCX 合并单元格会重复）
    col_count = max(len(r) for r in rows)
    md_rows = []
    for row in rows:
        # 补齐列数
        while len(row) < col_count:
            row.append('')
        md_rows.append('| ' + ' | '.join(row) + ' |')

    # 插入分隔行
    sep = '| ' + ' | '.join(['---'] * col_count) + ' |'
    result = [md_rows[0], sep] + md_rows[1:]
    return '\n'.join(result)


# ──────────────────────────────────────────────
# 5. 主处理流程
# ──────────────────────────────────────────────
def convert(input_path: str, output_path: str):
    doc = Document(input_path)

    # 收集段落和表格，按文档顺序处理
    # python-docx 的 doc.paragraphs 和 doc.tables 是独立列表
    # 用 element 树按顺序遍历
    from docx.oxml.ns import qn
    body = doc.element.body

    lines: list[str] = []
    table_index = 0
    para_index  = 0

    prev_line = ''

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':  # 段落
            if para_index < len(doc.paragraphs):
                para = doc.paragraphs[para_index]
                para_index += 1
            else:
                continue

            raw = para.text
            if is_noise(raw):
                continue

            text = fix_text(raw)
            if not text:
                continue
            # 二次噪声检测：fix_text 可能合并了空格后暴露新的噪声模式
            if is_noise(text):
                continue

            prev_empty = (prev_line == '')
            prefix, text = detect_structure(text, prev_empty)
            line = prefix + text

            # 避免连续重复行
            if line == prev_line:
                continue

            lines.append(line)
            prev_line = line

        elif tag == 'tbl':  # 表格
            if table_index < len(doc.tables):
                table = doc.tables[table_index]
                table_index += 1
            else:
                continue

            md_table = table_to_md(table)
            if md_table:
                lines.append('')
                lines.append(md_table)
                lines.append('')
                prev_line = ''

    # 后处理：合并多余空行，清理首尾
    output_lines = []
    blank_count  = 0
    for line in lines:
        if line == '':
            blank_count += 1
            if blank_count <= 1:
                output_lines.append(line)
        else:
            blank_count = 0
            output_lines.append(line)

    # 添加文档标题
    header = "# 种子课——一个数学特级教师的思与行\n\n**作者：俞正强**\n\n---\n"
    content = header + '\n'.join(output_lines).strip() + '\n'

    Path(output_path).write_text(content, encoding='utf-8')
    print(f"完成！输出文件：{output_path}")
    print(f"输入段落数：{len(doc.paragraphs)}")
    print(f"输出行数：{len(output_lines)}")

    # 统计
    h2 = sum(1 for l in output_lines if l.startswith('## '))
    h3 = sum(1 for l in output_lines if l.startswith('### '))
    tables_found = sum(1 for l in output_lines if l.startswith('|'))
    print(f"识别章节(##)：{h2}，识别小节(###)：{h3}，表格行数：{tables_found}")

    # 预览前50行
    print("\n=== 输出预览（前50行）===")
    for l in output_lines[:50]:
        print(repr(l) if not l else l)


if __name__ == '__main__':
    convert(INPUT, OUTPUT)
