# -*- coding: utf-8 -*-
"""
通用 OCR/扫描版 DOCX → Markdown 数据清洗脚本 v3
适用场景：任意扫描版/OCR识别的 Word 书籍
使用方法：只改下方「★ 每本书必须配置的区域 ★」，其余不用动

改进历史：
  v1 clean_docx.py：纯文字噪声过滤
  v2 clean_docx_v2.py：增加图片提取分类
  v3 (2026-03-17)：整合《低头找幸福》处理复盘的8条改进：
    - 字号优先识别标题（不再只靠文字正则）
    - OCR字号漂移：KNOWN_HEADINGS 固定文本字典
    - 噪声规则加入空格变体
    - 内置字号预分析（--analyze 参数）
    - 内置清洗后质量报告（标题数量检查）
    - 图片审查：自动生成 HTML 报告（--review-images 参数）
"""
import sys, re, os, zipfile, argparse
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from pathlib import Path
from lxml import etree

# ╔══════════════════════════════════════════════════════════════╗
# ║       ★ 每本书必须配置的区域 ★                               ║
# ║  步骤：先运行 --analyze 看字号分布，再填这里，再正式清洗       ║
# ╚══════════════════════════════════════════════════════════════╝

INPUT   = Path("/tmp/新课标数学课例式解读(OCR).docx")
OUT_DIR = Path("/tmp/新课标_cleaned")
BOOK_TITLE   = "新课标数学课例式解读（小学）"
BOOK_AUTHOR  = "孙晓天，张丹"

# 字号 → Markdown 标题级别映射
# 先运行 python clean_docx_v2.py --analyze 查看字号分布，再填写
# 不确定就留空（{}），脚本会退回文字正则识别
SIZE_TO_HEADING = {
    16.0: 2,   # "上编" / "下编" — 全书分部
    15.5: 2,   # "关键问题解读" — 主要章节
    14.5: 2,   # "推荐序" / "编者序" — 序言章节
    14.0: 3,   # "学术顾问" 及同级子节
    13.5: 3,   # "课例" 标题（大字号版）
    12.5: 4,   # "第一章" 及各章节
    12.0: 4,   # "课例" 标题（小字号版）
}

# 无论字号是多少，这些文本固定识别为 #### 标题（防OCR字号漂移）
KNOWN_HEADINGS_H4 = {
    '课例',
}

# 本书特有页眉/页脚：每页顶部都有"义务教育课程标准(2022年版) 课例式解读"
BOOK_HEADER_PATTERNS = [
    # 页眉：每页顶部的丛书名
    re.compile(r'^义务教育课程标准\s*[\(（]2022年版[\)）]'),
    re.compile(r'^小学数学\s*$'),
    re.compile(r'^课\s+例\s+式\s+解\s+读\s*$'),
    # 目录行：第X章/第X节/问题X 开头，末尾跟空白+页码数字（如 "第一章基于... 04"）
    re.compile(r'^(?:第[一二三四五六七八九十百]+[章节]|问题\d+)\s.{3,}[\s\t]\d{2,4}\s*$'),
]

# ╔══════════════════════════════════════════════════════════════╗
# ║       以下参数一般不需要改                                     ║
# ╚══════════════════════════════════════════════════════════════╝

OUT_MD  = OUT_DIR / f"{BOOK_TITLE}_cleaned.md"
IMG_DIR = OUT_DIR / "images"

SMALL_AREA   = 5000
WHITE_THRESH = 230

REL_NS  = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
A_BLIP  = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
VML_IMG = '{urn:schemas-microsoft-com:vml}imagedata'
R_EMBED = f'{{{REL_NS}}}embed'
R_ID    = f'{{{REL_NS}}}id'

# 这些字号确认为页眉/页脚/碎片，直接按字号过滤
NOISE_FONT_SIZES = {2.5, 3.0, 4.0, 5.0, 5.5, 6.0, 6.5, 7.0, 7.5, 8.0, 8.5, 9.0, 9.5}


# ════════════════════════════════════════════════════════════════
# 工具函数：字号获取
# ════════════════════════════════════════════════════════════════
def get_para_font_size(para):
    for run in para.runs:
        if run.font.size:
            return round(run.font.size.pt, 1)
    return None


# ════════════════════════════════════════════════════════════════
# 【--analyze 模式】字号分布分析（处理新书前必须先跑这个）
# ════════════════════════════════════════════════════════════════
def analyze_font_sizes():
    """打印字号分布，帮助配置 SIZE_TO_HEADING"""
    from collections import Counter
    doc = Document(INPUT)
    size_counts = Counter()
    size_samples = {}

    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        sz = get_para_font_size(p)
        bold = any(r.bold for r in p.runs if r.text)
        key = (sz, bold)
        size_counts[key] += 1
        if key not in size_samples:
            size_samples[key] = p.text.strip()[:60]

    print(f'\n{"="*70}')
    print(f'字号分析：{INPUT.name}')
    print(f'{"="*70}')
    print(f'{"字号":>8}  {"加粗":>4}  {"次数":>6}  示例文本')
    print('-'*70)
    for (sz, bold), cnt in sorted(size_counts.items(), key=lambda x: (-(x[0][0] or 0),)):
        sz_str = f'{sz}pt' if sz else 'None'
        flag = '★' if cnt <= 3 else ' '  # ★ 标记孤立字号（危险信号）
        print(f'{flag}{sz_str:>8}  {str(bold):>4}  {cnt:>6}  {size_samples[(sz,bold)]}')

    print(f'\n提示：')
    print(f'  ★ 标记的是出现次数 ≤3 的孤立字号，可能是 OCR 乱码，需人工确认')
    print(f'  根据上表填写脚本顶部的 SIZE_TO_HEADING 配置')
    print(f'  字号最大的几组通常是章节标题（##），中等是小节（###/####），正文最多')


# ════════════════════════════════════════════════════════════════
# 1. 噪声过滤规则
# ════════════════════════════════════════════════════════════════
_BASE_NOISE = [
    # 水印
    re.compile(r'[仅供个人科研教学使用！!]{4,}'),
    re.compile(r'供个人科研教学'),
    re.compile(r'人科研教学使用'),
    re.compile(r'更多教学好书请加微信'),
    re.compile(r'扫码关注'),
    re.compile(r'版权所有\s*[，,]?\s*翻印必究'),
    re.compile(r'未经许可.{0,10}不得'),
    # 纯页码
    re.compile(r'^\d{1,4}$'),
    # 全大写拼音
    re.compile(r'^[A-Z\s]{10,}$'),
    re.compile(r'^([A-Z]{2,}\s+){3,}'),
    # 出版/版权信息（含空格变体）
    re.compile(r'^(出版人|策划编辑|责任编辑|责任美编|封面设计|版式设计|责任校对|责任印制|项目统筹)'),
    re.compile(r'^(ISBN|CIP|图书在版编目|中国版本图书馆)'),
    re.compile(r'^(定价|传\s*真|邮\s*编|网\s*址|市场部电话|编辑部电话|社\s*址)'),
    re.compile(r'^(出版发行|出\s*版\s*发\s*行|印\s*刷|经\s*销|制\s*作|开\s*本|印\s*张|字\s*数|印\s*次|版\s*次|印\s*数)'),
    re.compile(r'(教育科学出版社|北京师范大学出版|人民教育出版社|华东师范大学出版)'),
    re.compile(r'(出版人|策划编辑).*(责任编辑|项目统筹)'),
    # 免责声明
    re.compile(r'如有印装质量问题'),
    re.compile(r'到所购图书销售'),
    # 丛书/系列
    re.compile(r'教育家[书書]院丛书'),
    # 孤立人名行（有内部空格）
    re.compile(r'^[\u4e00-\u9fff]{1,5}(\s+[\u4e00-\u9fff]{1,5}){1,5}\s*$'),
    re.compile(r'^[\u4e00-\u9fff]\s{1,3}[\u4e00-\u9fff]\s{1,3}[\u4e00-\u9fff]?\s*$'),
    # 目录行（含 /数字 页码）
    re.compile(r'.+/\d+\s*$'),
    # 目录标题行（"目 录 数字"）
    re.compile(r'^目\s+录\s+\d+\s*$'),
]

def build_noise_patterns():
    """合并基础噪声 + 本书特有页眉模式"""
    return _BASE_NOISE + BOOK_HEADER_PATTERNS

NOISE_PATTERNS = build_noise_patterns()
PUBLISHER_TABLE_KW = {'出版发行', '出 版发行', '出版 发行', '社址', '社   址',
                      '社  址', '社 址', '印刷', '经销', '制作',
                      '邮    编', '传   真', '网    址'}


def is_noise(text):
    t = text.strip()
    if not t:
        return True
    return any(p.search(t) for p in NOISE_PATTERNS)


# ════════════════════════════════════════════════════════════════
# 2. 文本修复 + 格式标准化
# ════════════════════════════════════════════════════════════════
def _is_spaced_chinese(text):
    chinese = re.findall(r'[\u4e00-\u9fff]', text)
    if len(chinese) < 4:
        return False
    ratio = len(chinese) / max(len(text.replace(' ', '')), 1)
    spaced = re.findall(r'[\u4e00-\u9fff] [\u4e00-\u9fff]', text)
    return ratio > 0.5 and len(spaced) >= 2


def fix_text(text):
    t = text.strip()
    t = re.sub(r'\s*[\|｜]\s*\d{1,4}\s*$', '', t)
    if _is_spaced_chinese(t):
        t = re.sub(r'(?<=[\u4e00-\u9fff，。！？、『』【】…])\s+(?=[\u4e00-\u9fff，。！？、『』【】…])', '', t)
    t = re.sub(r'[ \t]{2,}', ' ', t)
    if re.search(r'[\u4e00-\u9fff]', t):
        t = t.replace(',', '，').replace('?', '？').replace('!', '！')
        t = t.replace('(', '（').replace(')', '）')
    return t


# ════════════════════════════════════════════════════════════════
# 3. 标题检测（字号优先 → 固定文本 → 文字正则）
# ════════════════════════════════════════════════════════════════
H1_PAT = re.compile(r'^[一二三四五六七八九十]+[、.．]\s*.{2,30}$')
H3_PAT = re.compile(r'^[（\(][一二三四五六七八九\d]+[）\)]\s*.+$')


def detect_heading(text, font_size):
    t = text.strip()

    # 最高优先：固定文本标题（防OCR字号漂移）
    if t in KNOWN_HEADINGS_H4:
        return '#### ', t

    # 次优：字号映射
    if font_size and font_size in SIZE_TO_HEADING:
        level = SIZE_TO_HEADING[font_size]
        return '#' * level + ' ', t

    # 兜底：文字正则
    if H1_PAT.match(t):
        return '## ', t
    if H3_PAT.match(t):
        return '#### ', t
    if t.startswith('——') or t.startswith('— —'):
        return '> ', t
    if t.startswith('—') and len(t) < 40:
        return '> ', t

    return '', t


# ════════════════════════════════════════════════════════════════
# 4. 表格转 Markdown
# ════════════════════════════════════════════════════════════════
def table_to_md(table):
    rows = []
    for row in table.rows:
        cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
        rows.append(cells)
    rows = [r for r in rows if any(c for c in r)]
    if not rows:
        return ''
    first_cell = rows[0][0] if rows[0] else ''
    if any(kw in first_cell for kw in PUBLISHER_TABLE_KW):
        return ''
    all_text = [c for r in rows for c in r if c.strip()]
    if len(all_text) < 3:
        return ''
    col_count = max(len(r) for r in rows)
    md_rows = []
    for row in rows:
        while len(row) < col_count:
            row.append('')
        md_rows.append('| ' + ' | '.join(row) + ' |')
    sep = '| ' + ' | '.join(['---'] * col_count) + ' |'
    return '\n'.join([md_rows[0], sep] + md_rows[1:])


# ════════════════════════════════════════════════════════════════
# 5. 图片提取与分类
# ════════════════════════════════════════════════════════════════
def extract_and_classify_images(docx_path, images_dir):
    os.makedirs(images_dir, exist_ok=True)
    rid_to_path = {}
    stats = {'deleted': 0, 'kept': 0, 'blank': 0}

    with zipfile.ZipFile(docx_path) as z:
        rels_file = 'word/_rels/document.xml.rels'
        if rels_file not in z.namelist():
            return {}, stats
        with z.open(rels_file) as f:
            rels_tree = etree.parse(f)
        for rel in rels_tree.findall('{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
            if 'image' not in rel.get('Type', '').lower():
                continue
            target = rel.get('Target', '')
            rid = rel.get('Id', '')
            src = f'word/{target}' if not target.startswith('/') else target.lstrip('/')
            if src in z.namelist():
                ext = Path(target).suffix.lower()
                dst = os.path.join(images_dir, f'{rid}{ext}')
                with z.open(src) as img_src, open(dst, 'wb') as out:
                    out.write(img_src.read())
                rid_to_path[rid] = dst

    rid_to_md = {}
    try:
        from PIL import Image
        has_pil = True
    except ImportError:
        has_pil = False

    for rid, path in rid_to_path.items():
        fname = os.path.basename(path)
        rel_path = f'images/{fname}'
        if not has_pil:
            rid_to_md[rid] = f'\n![图片]({rel_path})\n'
            stats['kept'] += 1
            continue
        try:
            img = Image.open(path)
            w, h = img.size
            area = w * h
            ratio = max(w, h) / max(min(w, h), 1)
            gray = img.convert('L')
            pixels = list(gray.getdata())
            avg = sum(pixels) / len(pixels)
            img.close()
            if area < SMALL_AREA:
                os.remove(path); stats['deleted'] += 1; continue
            if ratio > 8:
                os.remove(path); stats['deleted'] += 1; continue
            if avg > WHITE_THRESH:
                os.remove(path); stats['blank'] += 1; continue
            if avg < 10 and area < 50000:
                os.remove(path); stats['deleted'] += 1; continue
            rid_to_md[rid] = f'\n![图片]({rel_path})\n'
            stats['kept'] += 1
        except Exception:
            if os.path.exists(path):
                os.remove(path)
            stats['deleted'] += 1

    return rid_to_md, stats


def get_para_image_rids(para_elem):
    rids = []
    for blip in para_elem.iter(A_BLIP):
        rid = blip.get(R_EMBED) or blip.get(R_ID)
        if rid:
            rids.append(rid)
    for imgdata in para_elem.iter(VML_IMG):
        rid = imgdata.get(R_ID) or imgdata.get(R_EMBED)
        if rid:
            rids.append(rid)
    return rids


# ════════════════════════════════════════════════════════════════
# 6. 图片审查 HTML 报告（--review-images 模式）
# ════════════════════════════════════════════════════════════════
def generate_image_review():
    """生成包含全部图片的 HTML 审查报告，方便人工确认哪些要保留"""
    import io
    try:
        from PIL import Image
    except ImportError:
        print('需要安装 Pillow：pip install Pillow')
        return

    review_dir = OUT_DIR / 'image_review'
    review_dir.mkdir(parents=True, exist_ok=True)
    kept_dir    = review_dir / '01_kept'
    small_dir   = review_dir / '02_deleted_small'
    line_dir    = review_dir / '03_deleted_lines'
    blank_dir   = review_dir / '04_deleted_blank'
    for d in [kept_dir, small_dir, line_dir, blank_dir]:
        d.mkdir(exist_ok=True)

    kept_rows = ''
    del_rows  = ''
    kept_n = del_n = 0

    with zipfile.ZipFile(INPUT) as z:
        with z.open('word/_rels/document.xml.rels') as f:
            rels_tree = etree.parse(f)
        for rel in rels_tree.findall('{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
            if 'image' not in rel.get('Type', '').lower():
                continue
            target = rel.get('Target', '')
            rid = rel.get('Id', '')
            src = f'word/{target}' if not target.startswith('/') else target.lstrip('/')
            if src not in z.namelist():
                continue
            data = z.open(src).read()
            ext  = Path(target).suffix.lower()
            fname = f'{rid}{ext}'
            try:
                img = Image.open(io.BytesIO(data))
                w, h = img.size
                area = w * h
                ratio = max(w, h) / max(min(w, h), 1)
                gray = img.convert('L')
                avg = sum(gray.getdata()) / (w * h)
                img.close()
                if area < SMALL_AREA:
                    reason = f'小碎片(面积{area}px²)'; folder = small_dir
                elif ratio > 8:
                    reason = f'分隔线(比例{round(ratio,1)})'; folder = line_dir
                elif avg > WHITE_THRESH:
                    reason = f'空白/水印(亮度{round(avg,1)})'; folder = blank_dir
                elif avg < 10 and area < 50000:
                    reason = f'纯黑装饰'; folder = small_dir
                else:
                    reason = f'{w}×{h} 亮度{round(avg,1)}'; folder = kept_dir
            except Exception as e:
                reason = f'打开失败:{e}'; folder = small_dir
                w = h = 0

            dst = folder / fname
            dst.write_bytes(data)
            rel_img = f'{folder.name}/{fname}'
            img_tag = f'<img src="{rel_img}" style="max-width:160px;max-height:160px;border:1px solid #ccc">'

            if folder == kept_dir:
                kept_rows += f'<tr><td style="font-size:11px">{rid}</td><td>{img_tag}</td><td style="color:green;font-size:11px">{reason}</td></tr>\n'
                kept_n += 1
            else:
                del_rows += f'<tr><td style="font-size:11px">{rid}</td><td>{img_tag}</td><td>{w}×{h}</td><td style="color:red;font-size:11px">{reason}</td></tr>\n'
                del_n += 1

    html = f"""<!DOCTYPE html><html lang="zh"><head><meta charset="utf-8">
<title>{BOOK_TITLE} 图片审查报告</title>
<style>body{{font-family:sans-serif;margin:20px}}table{{border-collapse:collapse;width:100%}}
td,th{{border:1px solid #ddd;padding:5px;vertical-align:middle}}th{{background:#f0f0f0}}</style>
</head><body>
<h1>《{BOOK_TITLE}》图片审查报告</h1>
<p style="background:#ffe;padding:10px;border:1px solid #cc0">
⚠️ <strong>必须先解压 zip，再双击此 HTML 文件打开</strong>，否则图片无法显示。</p>
<p>共 {kept_n+del_n} 张图片：<span style="color:green">保留 {kept_n} 张</span> / <span style="color:red">删除 {del_n} 张</span></p>
<h2>✅ 保留的 {kept_n} 张（自动判断有内容）</h2>
<table><tr><th>ID</th><th>预览</th><th>尺寸/亮度</th></tr>{kept_rows}</table>
<h2>🗑️ 删除的 {del_n} 张（请逐一确认是否误删）</h2>
<table><tr><th>ID</th><th>预览</th><th>尺寸</th><th>删除原因</th></tr>{del_rows}</table>
</body></html>"""

    (review_dir / '审查报告.html').write_text(html, encoding='utf-8')
    print(f'图片审查报告已生成：{review_dir / "审查报告.html"}')
    print(f'  保留: {kept_n}  删除: {del_n}')
    print(f'  解压后双击 审查报告.html 用浏览器打开查看')


# ════════════════════════════════════════════════════════════════
# 7. 主处理流程
# ════════════════════════════════════════════════════════════════
def convert():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    print('步骤 1/4：提取并分类图片...')
    rid_to_md, img_stats = extract_and_classify_images(INPUT, IMG_DIR)
    print(f'  保留: {img_stats["kept"]}  删除(碎片): {img_stats["deleted"]}  删除(空白): {img_stats["blank"]}')

    print('步骤 2/4：提取文本并清洗（字号优先识别标题）...')
    doc = Document(INPUT)
    body = doc.element.body
    lines = []
    prev_line = ''
    seen_lines = set()

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            para = Paragraph(child, doc)
            raw = para.text
            font_size = get_para_font_size(para)

            if font_size and font_size in NOISE_FONT_SIZES:
                for rid in get_para_image_rids(child):
                    if rid in rid_to_md:
                        lines.append(rid_to_md[rid])
                continue

            if is_noise(raw):
                for rid in get_para_image_rids(child):
                    if rid in rid_to_md:
                        lines.append(rid_to_md[rid])
                continue

            text = fix_text(raw)
            if not text or is_noise(text):
                continue

            text_key = text.strip().lower()
            if text_key in seen_lines and len(text_key) > 10:
                continue
            seen_lines.add(text_key)

            prefix, text = detect_heading(text, font_size)
            line = prefix + text
            if line == prev_line:
                continue

            lines.append(line)
            prev_line = line

            for rid in get_para_image_rids(child):
                if rid in rid_to_md:
                    lines.append(rid_to_md[rid])

        elif tag == 'tbl':
            table = Table(child, doc)
            md_table = table_to_md(table)
            if md_table:
                lines.append('')
                lines.append(md_table)
                lines.append('')
                prev_line = ''

    print('步骤 3/4：后处理...')
    output_lines = []
    blank_count = 0
    for line in lines:
        if line.strip() == '':
            blank_count += 1
            if blank_count <= 1:
                output_lines.append('')
        else:
            blank_count = 0
            output_lines.append(line)

    header = f'# {BOOK_TITLE}\n\n**作者：{BOOK_AUTHOR}**\n\n---\n'
    content = header + '\n'.join(output_lines).strip() + '\n'

    print('步骤 4/4：写入文件...')
    OUT_MD.write_text(content, encoding='utf-8')

    # ── 清洗质量报告 ──────────────────────────────────────────
    h2 = sum(1 for l in output_lines if l.startswith('## '))
    h3 = sum(1 for l in output_lines if l.startswith('### '))
    h4 = sum(1 for l in output_lines if l.startswith('#### '))
    img_refs = sum(1 for l in output_lines if '![图片]' in l)
    total_h = h2 + h3 + h4

    print(f'\n{"="*55}')
    print('数据清洗报告')
    print(f'{"="*55}')
    print(f'输入段落数:    {len(doc.paragraphs)}')
    print(f'输出行数:      {len(output_lines)}')
    print(f'## 章节:       {h2}')
    print(f'### 小节:      {h3}')
    print(f'#### 子节:     {h4}')
    print(f'标题合计:      {total_h}')
    print(f'图片引用:      {img_refs}')
    print(f'图片保留:      {img_stats["kept"]}')
    print(f'图片删除:      {img_stats["deleted"] + img_stats["blank"]}')
    print(f'输出文件:      {OUT_MD}')

    # 质量警告
    print(f'\n{"─"*55}')
    if total_h < 10:
        print(f'⚠️  警告：标题数量只有 {total_h} 个，可能 SIZE_TO_HEADING 未配置或配置错误')
        print(f'   建议：先运行 python {Path(__file__).name} --analyze 查看字号分布')
    else:
        print(f'✅ 标题数量正常（{total_h} 个）')
    if not SIZE_TO_HEADING:
        print(f'⚠️  警告：SIZE_TO_HEADING 为空，仅用文字正则识别标题，建议先做字号分析')


# ════════════════════════════════════════════════════════════════
# 入口
# ════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='OCR DOCX → Markdown 清洗工具')
    parser.add_argument('--analyze',       action='store_true', help='分析字号分布（处理新书前必做）')
    parser.add_argument('--review-images', action='store_true', help='生成图片审查 HTML 报告')
    args = parser.parse_args()

    if not INPUT.exists():
        print(f'输入文件不存在: {INPUT}')
        sys.exit(1)

    if args.analyze:
        analyze_font_sizes()
    elif args.review_images:
        OUT_DIR.mkdir(parents=True, exist_ok=True)
        generate_image_review()
    else:
        convert()
